#!/usr/bin/env python

#designed to operate on docx files only, for now

import sys
import re
from docx import Document

#globals
doc = Document(sys.argv[1])
foundSomething = False

#utility functions
def answerYes():
	ans = raw_input()
	if ans == 'y' or ans == 'Y' or ans == 'yes':
		return True
	elif ans == 'n' or ans == 'N' or ans == 'no':
		return False
	else:
		print "invalid answer..."
		return answerYes()

def isInt(c):
	try:
		int(c)
		return True
	except ValueError:
		return False

#arguments alternate between menu names and functions tied to those names
def menu(msg, *args):
	print '\n' + msg + '\n'
	count = 1
	for i in range(len(args)):
		if i % 2 == 0:
			print str(count) + ')  ' + args[i]
			count += 1
	choice = raw_input()
	if isInt(choice):
		choice = int(choice)
		if choice > 0 and choice <= len(args) / 2:
			args[choice * 2 - 1]()
			return

	print "invalid answer..."
	menu(msg, *args)

#rule tests
def paragraphLengthCheck():
	for ind, par in enumerate(doc.paragraphs):
		if len(par.text) > 500:
			print "Paragraph " + str(ind) + " (\"" + par.text[:20] + "...\") is " \
			+ str(len(par.text)) + " characters."
			foundSomething = True

def boldTimeStampsCheck():
	timeStampPattern = re.compile(r"[0-9][0-9]\:[0-9][0-9]\:[0-9][0-9]")
	timestamps = []
	for par in doc.paragraphs:
		timestamps += map(lambda x: x.group(), \
		re.finditer(timeStampPattern, par.text))

	print timestamps


def fullVerbatim():
	print "full Verbatim called"
	#paragraphLengthCheck()
	boldTimeStampsCheck()


def cleanVerbatim():
	print "clean verbatim called"
	paragraphLengthCheck()
	boldTimeStampsCheck()


#main
menu("Select a transcription type", "Full Verbatim", fullVerbatim, \
	"Clean Verbatim", cleanVerbatim)